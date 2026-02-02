#!/usr/bin/env python3
"""
Script untuk menambahkan Cover ke awal laporan dengan cara yang stabil.
Approach: Gunakan python-docx untuk copy paragraf dari cover.
"""
from pathlib import Path
from docx import Document
from docx.oxml import parse_xml
from copy import deepcopy

def add_cover_simple(cover_path, report_path, output_path=None):
    """
    Tambahkan cover ke awal laporan dengan cara yang stabil dan aman.
    """
    cover_path = Path(cover_path)
    report_path = Path(report_path)
    
    if output_path is None:
        output_path = report_path
    
    if not cover_path.exists():
        print(f"Error: Cover tidak ditemukan: {cover_path}")
        return False
    
    if not report_path.exists():
        print(f"Error: Report tidak ditemukan: {report_path}")
        return False
    
    print("Menambahkan Cover ke awal laporan (simple approach)...")
    
    # Load dokumen
    cover_doc = Document(cover_path)
    report_doc = Document(report_path)
    
    # Buat list dari paragraf cover
    cover_paragraphs = [p for p in cover_doc.paragraphs]
    
    print(f"✓ Ditemukan {len(cover_paragraphs)} paragraf di cover")
    
    # Insert paragraf cover ke awal report satu per satu
    report_body = report_doc.element.body
    
    # Salin XML elemen dari cover
    for i, cover_para in enumerate(cover_paragraphs):
        # Deep copy elemen XML
        para_copy = deepcopy(cover_para._element)
        # Insert di awal (index i)
        report_body.insert(i, para_copy)
    
    print(f"✓ {len(cover_paragraphs)} paragraf ditambahkan ke awal")
    
    # Insert page break setelah cover
    insert_page_break(report_doc, len(cover_paragraphs))
    
    # Simpan
    report_doc.save(output_path)
    print(f"✅ Cover berhasil ditambahkan: {output_path}")
    return True

def insert_page_break(doc, after_para_index):
    """Insert page break setelah paragraph index tertentu."""
    try:
        p = doc.paragraphs[after_para_index]
        pPr = p._element.get_or_add_pPr()
        pPgBr = parse_xml(r'<w:pageBreakBefore xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>')
        pPr.append(pPgBr)
        print("✓ Page break ditambahkan")
    except Exception as e:
        print(f"Warning: Page break gagal: {e}")

if __name__ == '__main__':
    root = Path(__file__).parent.parent / 'docs'
    cover_file = root / 'Cover.docx'
    report_file = root / 'laporan-proyek.docx'
    
    add_cover_simple(cover_file, report_file)
