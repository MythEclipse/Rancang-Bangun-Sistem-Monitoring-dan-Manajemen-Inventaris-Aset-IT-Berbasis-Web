#!/usr/bin/env python3
"""
Script untuk menambahkan warna pada heading di dokumen Word.
"""
from pathlib import Path
from docx import Document
from docx.shared import RGBColor

def add_heading_colors(docx_path):
    """
    Ubah warna heading menjadi hitam (auto):
    - Heading 1: Black (0, 0, 0)
    - Heading 2: Black (0, 0, 0)
    - Heading 3: Black (0, 0, 0)
    """
    docx_path = Path(docx_path)
    
    if not docx_path.exists():
        print(f"Error: File not found: {docx_path}")
        return False
    
    print("Mengubah warna heading menjadi hitam (auto)...")
    
    doc = Document(docx_path)
    
    # Warna untuk heading - semua hitam (auto)
    colors = {
        'Heading 1': RGBColor(0, 0, 0),    # Black (Auto)
        'Heading 2': RGBColor(0, 0, 0),    # Black (Auto)
        'Heading 3': RGBColor(0, 0, 0),    # Black (Auto)
    }
    
    # Apply warna ke paragraf dengan style heading
    for para in doc.paragraphs:
        style_name = para.style.name
        if style_name in colors:
            for run in para.runs:
                run.font.color.rgb = colors[style_name]
                print(f"✓ Warna ditambahkan ke {style_name}: {run.text[:50]}")
    
    # Apply warna ke table cells
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    style_name = para.style.name
                    if style_name in colors:
                        for run in para.runs:
                            run.font.color.rgb = colors[style_name]
    
    doc.save(docx_path)
    print(f"\n✅ Warna heading berhasil diubah: {docx_path}")
    print("   - Heading 1: Black (Auto)")
    print("   - Heading 2: Black (Auto)")
    print("   - Heading 3: Black (Auto)")
    return True

if __name__ == '__main__':
    docx_file = Path(__file__).parent.parent / 'docs' / 'laporan-proyek.docx'
    add_heading_colors(docx_file)
